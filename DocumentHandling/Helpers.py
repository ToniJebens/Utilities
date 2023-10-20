import os
import json
import time
import requests
import pandas as pd
import shutil
from PyPDF2 import PdfFileReader
import PyPDF2
from docx import Document
import docx


def get_doc_type(file_path):
    """
    Determine the type of a document based on its file extension.
    """
    _, file_extension = os.path.splitext(file_path)
    if file_extension in ['.txt']:
        return 'txt'
    elif file_extension in ['.pdf']:
        return 'pdf'
    elif file_extension in ['.docx']:
        return 'docx'
    else:
        print(f'Unsupported file extension: {file_extension}')
        return None


def get_file_creation_time(file_path: str) -> str:
    """
    Get the creation time of a file.

    :param file_path: path to the file
    :return: creation time of the file
    """
    if os.path.exists(file_path):
        creation_time_epoch = os.path.getctime(file_path)

        # convert to human-readable format
        creation_time = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(creation_time_epoch))
        return creation_time
    
    return "File not found."



def create_doc(doc_type, file_path, content):
    if doc_type == 'txt':
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
    elif doc_type == 'pdf':
        print("Writing PDF files is not implemented.")
    elif doc_type == 'docx':
        doc = docx.Document()
        doc.add_paragraph(content)
        doc.save(file_path)
    else:
        print(f"Cannot create document with extension {doc_type}.")
    return file_path



def load_pdf(file_path: str) -> str:
    if not os.path.isfile(file_path):
        raise ValueError(f'File not found: {file_path}')
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfFileReader(file)
            pdf = ' '.join([reader.getPage(i).extractText() for i in range(reader.numPages())])
            return pdf
    except Exception as e:
        print(f'Error while loading PDF: {e}')
        return None


def load_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding = 'utf-8') as f:
        return f.read()
    

def load_docx(file_path: str) -> str:
    doc = Document(file_path)
    return ' '.join([paragraph.text for paragraph in doc.paragraphs])


def load_document(fileExtension, file_path):
    if fileExtension == "pdf":
        doc = load_pdf(file_path)
    elif fileExtension == "txt":
        doc = load_txt(file_path)
    elif fileExtension in ["doc", "docx"]:
        doc = load_docx(file_path)
    else:
        print(f"Cannot read document with extension {fileExtension}.")
